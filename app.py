import streamlit as st
from google import genai
from docx import Document
import json
from answers import formAnswer
from questions import form_data
import os


client= genai.Client(api_key=st.secrets["API_KEY_2"])

def process_item(item):
    if isinstance(item, dict):
        temp_dict = {}
        for key, value in item.items():
            # If key starts with '{{', keep it and set value to 'answer'
            if key.startswith('{{'):
                temp_dict[key] = value
            # If the value is a dictionary, process it recursively
            elif isinstance(value, dict):
                processed = process_item(value)
                # Add all items from processed dictionary directly to the result
                for k, v in processed.items():
                    temp_dict[k] = v
        return temp_dict
    return item
#gemini-3-flash-preview
#gemini-2.5-flash
#gemini-2.5-flash-lite
def get_gemini_response(prompt):
    response = client.models.generate_content(model="gemini-2.5-flash-lite", contents=prompt)
    return response.text

def find_missing_keys(form1, form2):
    return list(set(form1.keys()) - set(form2.keys()))

APP_PASSWORD = st.secrets["PASS"]
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if not st.session_state.authenticated:
    st.title("🔒 Enter Password")
    password = st.text_input("Password:", type="password")
    if st.button("Login"):
        if password == APP_PASSWORD:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password. Try again.")
    st.stop()


st.title("Form Filler")
selected_form = st.selectbox("Select a Form:", list(form_data.keys()), index=0)
user_input = st.text_area("Enter The Prompt:", height=250)
response = ""
if st.button("Fill Form"):
    if user_input.strip():
        with st.spinner("Preparing your file..."):
            if(selected_form=="Form3"):

                query1 = "For the given prompt:"+user_input+",analyze it very well and understand it and then check these questions of the following dictionary:"+json.dumps(form_data.get(selected_form).get("Form3_part1"))+". I want the answers in the form of:"+json.dumps(formAnswer.get("Form3_part1"))+"and the answer of {{checks}} is either 'true' or 'false' only, and as a string only, if you do not know if it is 'true' or 'false' then put 'false', and analyze the prompt carefully before answering. Make sure to fill ALL the fields of the given sample and NEVER put a null value, put NA if and only if the answer can not be determined or concluded or interpreted or analyzed from the prompt, but avoid putting NA as much as possible. Give me the result directly in json format with nothing written before or after and DO NOT SKIP ANY ENTRY IN THE ANSWERS file or the dictionary file especially {{check_33}}, and the answers must not exceed 20 words, no 'answer' having more than 20 words is acceptable. Try to elaborate your answers within those 20 words even if there is no enough data try to analyze them, do not write very brief answers. Give me the result directly in json format as a 'String', with nothing written before or after! remember that the questions dictionary is not the one I want to fill, I want to fill the answers dictionary which is the 2nd one that I wrote"
                query2= "For the given prompt:"+user_input+",analyze it very well and understand it and then answer the questions of the following dictionary:"+json.dumps(form_data.get(selected_form).get("Form3_part2"))+", and fill this dictionary with the correct answers:"+json.dumps(formAnswer.get("Form3_part2"))+". Analyze the prompt carefully before answering. Make sure to fill ALL the fields of the given sample and NEVER put a null value, put NA if and only if the answer can not be determined or concluded or interpreted or analyzed from the prompt, but avoid putting NA as much as possible. Give me the result directly in json format with nothing written before or after and DO NOT SKIP ANY ENTRY IN THE ANSWERS file or the dictionary file especially {{check_33}}, and the answers must not exceed 20 words, no 'answer' having more than 20 words is acceptable. Try to elaborate your answers within those 20 words even if there is no enough data try to analyze them, do not write very brief answers. Give me the result directly in json format as a 'String', with nothing written before or after!"
                response1 = get_gemini_response(query1)
                st.write("Processing 1")
                if(response1[:3]=="```"):
                    response1= response1[8:len(response1)-4]
                ############################################
                if(response1[:3]=="\"{{"):
                    response1="{"+response1
                if(response1[-1]=='"'):
                    response1=response1+"}"
                response2=  get_gemini_response(query2)
                if(response2[:3]=="```"):
                    response2= response2[8:len(response2)-4]
                st.write("Processing 2")
                if(response2[:3]=="\"{{"):
                    response2="{"+response2
                if(response2[-1]=='"'):
                    response2=response2+"}"
                response = response1.rstrip('}') + ',' + response2.lstrip('{')
                st.write("Processing 3")
                st.write(response)
                response=process_item(response)
                st.write("Processing 4")
            elif(selected_form=="Form2"):
                st.write("Started Processing")
                query1 = "For the given prompt:"+user_input+",analyze it very well and understand it and then answer the questions of the following dictionary:"+json.dumps(form_data.get(selected_form).get("Form2_part1"))+", and fill this dictionary with the correct answers:"+json.dumps(formAnswer.get(selected_form))+"and the answer of {{checks}} is either 'true' or 'false' only, and as a string only, if you do not know if it is 'true' or 'false' then put 'false', and analyze the prompt carefully before answering. Make sure to fill ALL the fields of the given sample especially long questions and be strict to the word limit. Give me the result directly in json format with nothing written before or after and DO NOT SKIP ANY ENTRY IN THE ANSWERS file or the dictionary file especially {{check_33}}, and the answers must not exceed 100 words, no 'answer' having more than 100 words is acceptable, and answer directly without writing 'Person said that he/she'. Try to elaborate your answers within those 100 words even if there is no enough data try to analyze them and be reasonable, do not write brief answers. Give me the result directly in json format with nothing written before or after!"
                query2 = "For the given prompt:"+user_input+",analyze it very well and understand it and then answer the questions of the following dictionary:"+json.dumps(form_data.get(selected_form).get("Form2_part2"))+", and fill this dictionary with the correct answers:"+json.dumps(formAnswer.get("Form2_part2"))+"and the answer of {{checks}} is either 'true' or 'false' only, and as a string only, if you do not know if it is 'true' or 'false' then put 'false', and analyze the prompt carefully before answering. Make sure to fill ALL the fields of the given sample especially long questions and be strict to the word limit. Give me the result directly in json format with nothing written before or after and DO NOT SKIP ANY ENTRY IN THE ANSWERS file or the dictionary file especially {{check_33}}, and the answers must not exceed 100 words, no 'answer' having more than 100 words is acceptable, and answer directly without writing 'Person said that he/she'. Try to elaborate your answers within those 100 words even if there is no enough data try to analyze them and be reasonable, do not write brief answers. Give me the result directly in json format with nothing written before or after! Do not write the question itself in the json file, just the placeholder:value of each placeholder in order to use them to fill a form. Don't forget the {{violation_type}} entry, it is just after {{explain_6}}"
                response1 = get_gemini_response(query1)
                response2=  get_gemini_response(query2)
                st.write("Processing 1")
                #st.write(response1)
                #st.write(response2)
                print("______________________")
                print(response1)
                print("______________________")
                print(response2)
                print("______________________")
                if(response1[:3]=="```"):
                    response1= response1[8:len(response1)-4]
                    response2= response2[8:len(response2)-4]
                response = response1.rstrip('}') + ',' + response2.lstrip('{')
                print(response)
                st.write("Processing 2")
                response=process_item(response)
                st.write("Processing 3") 
                print("_______________________")
                print(response)
                #st.write(response)
            elif(selected_form=="CFP"):
                query = "For the given prompt:"+user_input+",analyze it very well and understand it and then answer the questions of the following dictionary:"+json.dumps(form_data.get(selected_form))+", and fill this dictionary with the correct answers:"+json.dumps(formAnswer.get(selected_form))+" and analyze the prompt carefully before answering. Make sure to fill ALL the fields especially long questions and be strict to the word limit. Give me the result directly in json format with nothing written before or after and DO NOT SKIP ANY ENTRY IN THE ANSWERS file or the dictionary file, and answer directly without writing 'Person said that he/she', and do not add any text formatting like bold or italic or anything else just a pure text. Give me the answers in the shape of a paragraph without a title or a conclusion or extra stuff. Try to elaborate your answers within the word limit (4000-5000 characters), even if there is no enough data try to analyze them and be reasonable, do not write brief answers. {{response_1}} must be 2500-3500 characters length, while the others can be 1500-2500 characters length only, no more than 3000 characters. Give me the result directly in json format with nothing written before or after!  be formal and answer as a 1 block paragraph for each and 1000-2000 characters for each {{response}}, so the total number of characters must be 5000-6500 characters, add details as much as you can. Give me the result directly in json format as a string with nothing written before or after! I want the full response combined to be of 5500-7000 length only, you are usually writing 8000 or 4000, which is not acceptable."
                response = get_gemini_response(query)
                response = response[8:len(response)-4]
                
                
                if(response[:3]=="\"{{"):
                    response="{"+response
                if(response[-1]=='"'):
                    response=response+"}"
                if(response[0]!='{' and response[-1]!='}'):
                    response="{"+response+"}"
                response=process_item(response)
    else:
        st.warning("Please enter a valid prompt!")

word_docs = {
    "Form1": "form1.docx",
    "Form2": "form2.docx",
    "Form3": "form3.docx",
    "Form4": "form4.docx",
    "Form5": "form5.docx",
    "Form6": "form6.docx",
    "CFP": "CFP.docx"
}
print(response, "\n\n",type(response))
doc = Document(word_docs.get(selected_form))

data = {}
# Dictionary of values to replace placeholders
if response != "":
    try:
        print(response[8:len(response)-4])
        st.write("filling 1")
        #st.write(response)
        data = json.loads(response)
        st.write("filling 2")
        data=process_item(data)
        st.write("filling 3")
        if(selected_form=="Form3"):
            all_key_values= formAnswer.get("Form3_part1")
            missing_keys = find_missing_keys(formAnswer.get("Form3_part1"), formAnswer.get("Form3_part2"))
            for key in missing_keys:
                if str(key).startswith("{{Check"):
                    all_key_values[key] = " ☐"
                else:
                    all_key_values[key] = " "
            missing_keys = find_missing_keys(all_key_values, data)
            for key in missing_keys:
                if str(key).startswith("{{Check"):
                    data[key] = " ☐"
                else:
                    data[key] = " "

        elif(selected_form=="Form2"):
            all_key_values= formAnswer.get("Form2")
            missing_keys = find_missing_keys(formAnswer.get("Form2"), formAnswer.get("Form2_part2"))
            for key in missing_keys:
                if str(key).startswith("{{Check"):
                    all_key_values[key] = " ☐"
                else:
                    all_key_values[key] = " "
            missing_keys = find_missing_keys(all_key_values, data)
            for key in missing_keys:
                if str(key).startswith("{{Check"):
                    data[key] = " ☐"
                else:
                    data[key] = " "
            st.write("filling 4")
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key in cell.text:
                            if str(value) == "false" or str(value) == "False" or str(value) == "No" or str(value) == "no":
                                if(str(key).startswith("{{Check")):
                                    cell.text = cell.text.replace(key, " ☐")
                                else:
                                    cell.text = cell.text.replace(key, value)
                            elif str(value) == "true" or str(value) == "True" or str(value) == "Yes" or str(value) == "yes":
                                if(str(key).startswith("{{Check")):
                                    cell.text = cell.text.replace(key, " ☑")
                                else:
                                    cell.text = cell.text.replace(key, value)
                            elif (value=="NA" or value=="answer" or value=="N/A" or value=="null" or value=="Null" or value==""):
                                if(str(key).startswith("{{Check")):
                                    cell.text = cell.text.replace(key, " ☐")
                                else:
                                    cell.text = cell.text.replace(key, " ")
                            else:
                                cell.text = cell.text.replace(str(key or ""), str(value or ""))
        st.write("filling 5")
        file_path = f"{selected_form}_filled.docx"
        filled = doc.save(file_path)
        st.write("✅ The form is successfully filled!")
        with open(file_path, "rb") as file:
            btn = st.download_button(
                label="📄 Download Filled Word Document",
                data=file,
                file_name=file_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Delete file after download button is clicked
        if btn:
            os.remove(file_path)

        print("✅ The form is successfully filled and deleted after downloading!")

    except json.JSONDecodeError:
        st.error("Try Again")


st.write(" ")
st.write(" ")
st.title("Case Notes Filler")


caseNoteDict = {
    "{time1}": "answer",
    "{first}": "answer",
    "{time2}": "answer",
    "{second}": "answer",
    "{time3}": "answer",
    "{third}": "answer",
    "{time4}": "answer",
    "{fourth}": "answer",
    "{time5}": "answer",
    "{fifth}": "answer",
    "{time6}": "answer",
    "{sixth}": "answer",
    "{time7}": "answer",
    "{seventh}": "answer",
    "{time8}": "answer",
    "{eighth}": "answer",
    "{time9}": "answer",
    "{ninth}": "answer",
    "{time10}": "answer",
    "{tenth}": "answer"
}

user_input2 = st.text_area("Enter:", height=250)
caseNoteResponse = ""
sampleAnswer = {
    "{time1}": "30 Nov",
    "{first}": "Phone Call (Follow-up on Assessment): This phone call was made to follow up on an earlier assessment. The objective was likely to check if the assessment results were clear or if any additional information or clarification was needed. The phone call may have involved discussing the outcomes of the assessment, addressing any concerns, and providing further details or instructions based on the findings.",
    
    "{time2}": "7 Dec",
    "{second}": "Home Visit (Action Plan and Provide Awareness Session PSs): During this visit, an action plan was presented and discussed, with the aim to outline steps or recommendations for improvement or progress. The awareness session, likely focused on providing key information or instructions, was conducted. The 'PSs' might refer to 'project steps,' 'problem-solving strategies,' or specific recommendations tied to the action plan, helping the person or group to understand and implement the next steps.",
    
    "{time3}": "20 Dec",
    "{third}": "Home Visit and Distribution of CRls: The visit was aimed at addressing further needs or providing updates on progress from the previous visit. 'CRls' likely refers to 'certificates of recognition,' 'checklist review logs,' or some kind of related materials being distributed. The home visit was to personally deliver these materials and ensure proper understanding of their contents.",
    
    "{time4}": "4 Jan",
    "{fourth}": "Phone Call: This phone call was likely made to check in after the previous interactions, offering further support, clarification, or an update. The call could have focused on following up on any changes, challenges, or progress since the last in-person meeting or distribution of materials.",
    
    "{time5}": "22 Jan",
    "{fifth}": "Phone Call and Follow-up: This call was made to follow up on ongoing matters, ensuring that previous recommendations or actions were being implemented and addressing any new developments or issues. The follow-up likely included reviewing progress, providing additional guidance if necessary, and possibly setting the stage for future steps or meetings.",
    
    "{time6}": "",
    "{sixth}": "",
    
    "{time7}": "",
    "{seventh}": "",
    
    "{time8}": "",
    "{eighth}": "",
    
    "{time9}": "",
    "{ninth}": "",
    
    "{time10}": "",
    "{tenth}": ""
}

caseNote = Document("CaseNotes.docx")
caseData = {}
if st.button("Fill Case Note"):
    if user_input2.strip():
        with st.spinner("Preparing your file..."):
            caseQuery = "For each of the following timelines " + user_input2 + "provide your own extremely detailed description and explanation of each and fill your explanation and description along with the timelines in the values of this dictionary " + json.dumps(caseNoteDict) + "by replacing the dictionary values as appropriate. Give me the dictionary directly with nothing before or after. For example, if my input is 30 Nov Phone call -follow up on assessment 7 dec home visit (action plan and provide awareness session PSs) 20 Dec home visit and distribution CRls 4 Jan phone call 22-jan phone call and follow up your detailed explanation could be "+ json.dumps(sampleAnswer) + " like this sample answer please"
            caseNoteResponse = get_gemini_response(caseQuery)
            print(caseNoteResponse, type(caseNoteResponse))
            caseData = json.loads(caseNoteResponse[8:len(caseNoteResponse)-4])
            for table in caseNote.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in caseData.items():
                            if key in cell.text:
                                if not value:
                                    cell.text = cell.text.replace(str(key), "")
                                else:
                                    cell.text = cell.text.replace(str(key), str(value))
                            
            file_path2 = "CaseNotes_filled.docx"
            filled2 = caseNote.save(file_path2)
            st.write("✅ The form is successfully filled!")
            with open(file_path2, "rb") as file:
                btn = st.download_button(
                label="📄 Download Filled Case NotesDocument",
                data=file,
                file_name=file_path2,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Delete file after download button is clicked
        if btn:
            os.remove(file_path2)

        print("✅ The form is successfully filled and deleted after downloading!")
    else:
        st.warning("Please enter a valid prompt!")

































